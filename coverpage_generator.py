import re

from docx import Document


async def paragraph_replace_text(paragraph, regex, replace_str):
    """Return `paragraph` after replacing all matches for `regex` with `replace_str`.

    `regex` is a compiled regular expression prepared with `re.compile(pattern)`
    according to the Python library documentation for the `re` module.
    """
    # --- a paragraph may contain more than one match, loop until all are replaced ---
    while True:
        text = paragraph.text
        match = regex.search(text)
        if not match:
            break

        # --- when there's a match, we need to modify run.text for each run that
        # --- contains any part of the match-string.
        runs = iter(paragraph.runs)
        start, end = match.start(), match.end()

        # --- Skip over any leading runs that do not contain the match ---
        for run in runs:
            run_len = len(run.text)
            if start < run_len:
                break
            start, end = start - run_len, end - run_len

        # --- Match starts somewhere in the current run. Replace match-str prefix
        # --- occurring in this run with entire replacement str.
        run_text = run.text
        run_len = len(run_text)
        run.text = "%s%s%s" % (run_text[:start], replace_str, run_text[end:])
        end -= run_len  # --- note this is run-len before replacement ---

        # --- Remove any suffix of match word that occurs in following runs. Note that
        # --- such a suffix will always begin at the first character of the run. Also
        # --- note a suffix can span one or more entire following runs.
        for run in runs:  # --- next and remaining runs, uses same iterator ---
            if end <= 0:
                break
            run_text = run.text
            run_len = len(run_text)
            run.text = run_text[end:]
            end -= run_len

    # --- optionally get rid of any "spanned" runs that are now empty. This
    # --- could potentially delete things like inline pictures, so use your judgement.
    # for run in paragraph.runs:
    #     if run.text == "":
    #         r = run._r
    #         r.getparent().remove(r)

    return paragraph

async def document_replace_text(document, regex, replace_str):
    """Return `document` after replacing all matches for `regex` with `replace_str`.

    `regex` is a compiled regular expression prepared with `re.compile(pattern)`
    according to the Python library documentation for the `re` module.
    """
    for paragraph in document.paragraphs:
        paragraph_replace_text(paragraph, regex, replace_str)

    for table in document.tables:
        for cell in table._cells:
            for paragraph in cell.paragraphs:
                paragraph_replace_text(paragraph, regex, replace_str)
    return document


if __name__ == "__main__":

    document = Document("./Lab Report Template.docx")

    regex1 = re.compile("<course title>")
    regex2 = re.compile("<course code>")
    regex3 = re.compile("<id>")
    regex4 = re.compile("<teacher name>")
    regex5 = re.compile("<teacher designation>")
    regex6 = re.compile("<date>")

    replace_str1 = "Pharmacology Lab - II"
    replace_str2 = "PHR-302"
    replace_str3 = "18PHR053"
    replace_str4 = "Dr. S. K. Saha"
    replace_str5 = "Assistant Professor"
    replace_str6 = "01/01/2018"


    document_replace_text(document, regex1, replace_str1)
    document_replace_text(document, regex2, replace_str2)
    document_replace_text(document, regex3, replace_str3)
    document_replace_text(document, regex4, replace_str4)
    document_replace_text(document, regex5, replace_str5)
    document_replace_text(document, regex6, replace_str6)
    document.save('./Lab Report Coverpage.docx')
